from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

BASE = Path(__file__).resolve().parent
OUT = BASE / "MarginLift-x-Chargoon-Proposal-fa.docx"
LOGO = BASE / "assets" / "logo" / "Fa" / "logo-01.png"
SCREEN = BASE.parent.parent / "qa-retention-analysis-desktop.png"

FONT = "Vazirmatn"
INK = "111A16"
GREEN = "046A38"
MID = "43B02A"
LIME = "C4D600"
MUTED = "627068"
SOFT = "EAF6EE"
PAPER = "F6F8F5"
LINE = "D9E1DC"
AMBER = "B36B00"
RED = "B3312C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"] + (["insideH", "insideV"] if inside else [])
    for edge in edges:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_rtl(table):
    tbl_pr = table._tbl.tblPr
    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    return paragraph


def para(doc_or_cell, text="", size=10.5, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT,
         before=0, after=6, line=1.25, keep=False, italic=False):
    p = doc_or_cell.add_paragraph()
    rtl(p, align)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if text:
        set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    rtl(p)
    p.paragraph_format.page_break_before = False
    set_run(p.add_run(text), size=19, color=GREEN, bold=True)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    rtl(p)
    set_run(p.add_run(text), size=13.5, color=INK, bold=True)
    return p


def label(doc, text, color=GREEN):
    p = para(doc, text, size=8.5, color=color, bold=True, after=2, keep=True)
    return p


def add_page_num(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run(run, size=8, color=MUTED)


def page_break(doc):
    doc.add_page_break()


def callout(doc, title, body, fill=SOFT, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.65)
    set_table_borders(table, color=fill, size=0, inside=False)
    set_table_rtl(table)
    cell = table.cell(0, 0)
    cell.width = Inches(6.65)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=170, start=220, bottom=170, end=220)
    p1 = cell.paragraphs[0]
    rtl(p1)
    p1.paragraph_format.space_after = Pt(4)
    set_run(p1.add_run(title), size=10, color=accent, bold=True)
    body_color = WHITE if fill == INK else INK
    para(cell, body, size=12.2, color=body_color, bold=True, after=0, line=1.25)
    para(doc, "", size=2, after=2)
    return table


def card_grid(doc, cards, cols=2):
    rows = (len(cards) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_rtl(table)
    set_table_borders(table, color=LINE, size=6, inside=True)
    widths = [Inches(3.30)] * cols
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            cell.width = widths[c]
            set_cell_margins(cell, top=170, start=180, bottom=170, end=180)
            idx = r * cols + c
            if idx >= len(cards):
                set_cell_shading(cell, WHITE)
                continue
            title, body, highlight = cards[idx]
            set_cell_shading(cell, INK if highlight else WHITE)
            p = cell.paragraphs[0]
            rtl(p)
            p.paragraph_format.space_after = Pt(5)
            set_run(p.add_run(title), size=12.5, color=WHITE if highlight else GREEN, bold=True)
            para(cell, body, size=9.6, color="DCE5E0" if highlight else MUTED, after=0, line=1.18)
    para(doc, "", size=2, after=2)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    rtl(p)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.24)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    set_run(p.add_run(text), size=10.2, color=INK)
    return p


def set_fixed_table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.inches * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(17)
section.bottom_margin = Mm(16)
section.left_margin = Mm(17)
section.right_margin = Mm(17)
section.header_distance = Mm(7)
section.footer_distance = Mm(7)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.25
for style_name, size, color, before, after in [
    ("Heading 1", 19, GREEN, 14, 7),
    ("Heading 2", 13.5, INK, 10, 5),
    ("Heading 3", 11.5, GREEN, 8, 4),
]:
    st = styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st._element.rPr.rFonts.set(qn("w:cs"), FONT)
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
rtl(hp)
set_run(hp.add_run("پیشنهاد همکاری مارجین‌لیفت × چارگون"), size=8, color=MUTED, bold=True)
footer = section.footer
fp = footer.paragraphs[0]
add_page_num(fp)
set_run(fp.add_run("   |   محرمانه و غیرالزام‌آور"), size=8, color=MUTED)

# Cover
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_rtl(table)
set_table_borders(table, color=WHITE, size=0, inside=False)
set_fixed_table_widths(table, [Inches(3.25), Inches(3.40)])
left, right = table.rows[0].cells
left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
left_p = left.paragraphs[0]
left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(left_p.add_run("MARGINLIFT"), size=13, color=GREEN, bold=True)
right_p = right.paragraphs[0]
right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
right_p.add_run().add_picture(str(LOGO), width=Inches(1.75))

para(doc, "", size=4, after=44)
label(doc, "پیشنهاد همکاری محصولی", GREEN)
para(doc, "از داده سازمانی\nتا تصمیم قابل‌اندازه‌گیری", size=28, color=INK, bold=True, after=12, line=1.10)
para(doc, "لایه تصمیم علّی برای نگهداشت مشتری و تخصیص هوشمند مشوق؛ مکمل محصولات هوش مصنوعی، هوش تجاری و تعامل‌پذیری چارگون", size=13, color=MUTED, after=22, line=1.35)
callout(doc, "تز همکاری", "چارگون داده و جریان کار سازمانی را یکپارچه می‌کند؛ مارجین‌لیفت این جریان را به اقدام و نتیجه مالی قابل‌سنجش متصل می‌کند.", fill=INK, accent=LIME)
para(doc, "", size=4, after=25)
meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(meta)
set_table_borders(meta, color=LINE, size=5, inside=True)
set_fixed_table_widths(meta, [Inches(3.32), Inches(3.32)])
metadata = [
    ("مخاطب", "جناب آقای شاهین طبری\nبنیان‌گذار و رئیس هیئت‌مدیره چارگون"),
    ("نوع سند", "پیشنهاد اولیه همکاری\nغیرالزام‌آور و محرمانه"),
    ("تهیه‌کننده", "تیم مارجین‌لیفت\nmarginlift.ir"),
]
for r, (label_text, value_text) in enumerate(metadata):
    for c, text in enumerate((value_text, label_text)):
        cell = meta.cell(r, c)
        set_cell_shading(cell, WHITE if c == 0 else SOFT)
        set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
        p = cell.paragraphs[0]
        rtl(p)
        set_run(p.add_run(text), size=9.5 if c == 0 else 9, color=INK if c == 0 else GREEN, bold=c == 1)
para(doc, "مرداد ۱۴۰۵", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=0)

# 2 — Executive letter
page_break(doc)
label(doc, "خطاب مدیریتی")
h1(doc, "پیشنهاد در یک نگاه")
para(doc, "جناب آقای طبری،", size=11.5, bold=True, after=7)
para(doc, "چارگون طی بیش از دو دهه، پیچیدگی نرم‌افزار سازمانی را به سامانه‌های یکپارچه، امن و قابل اتکا تبدیل کرده است. مسیر جدید این شرکت نیز روشن است: تجربه انسانی‌تر، هوش مصنوعی درون‌سازمانی، تعامل‌پذیری و همکاری با محصولات مکمل.", size=11, after=9, line=1.35)
para(doc, "پیشنهاد مارجین‌لیفت افزودن یک داشبورد دیگر نیست. پیشنهاد ما، یک لایه تخصصی برای تصمیم‌های نگهداشت مشتری است: چه کسی در معرض ریزش است، برای چه کسی اقدام واقعاً رفتار را تغییر می‌دهد، چه نوع اقدام یا مشوقی اقتصادی است و نتیجه چگونه با گروه کنترل سنجیده می‌شود.", size=11, after=9, line=1.35)
callout(doc, "پیشنهاد تصمیم", "یک پایلوت هم‌برند با یک مشتری مناسب چارگون؛ ابتدا ارزیابی داده، سپس اجرای کنترل‌شده و در پایان تصمیم شفاف برای توسعه، بازبینی یا توقف.", fill=SOFT, accent=GREEN)
h2(doc, "آنچه در این مرحله درخواست می‌کنیم")
add_bullet(doc, "یک جلسه ۶۰ دقیقه‌ای با حضور نمایندگان محصول، آسا، هوش تجاری، تعامل‌پذیری و امنیت.")
add_bullet(doc, "انتخاب یک مشتری با خرید تکرارشونده، هزینه اقدام مشخص و مالک KPI روشن.")
add_bullet(doc, "توافق روی قرارداد حداقلی داده و معیار موفقیت پیش از هرگونه توسعه اختصاصی.")
para(doc, "این مسیر عمداً کم‌ریسک طراحی شده است: تا زمانی که داده و نتیجه واقعی، ارزش را اثبات نکنند، پیشنهاد توسعه ماژول یا همکاری تجاری بلندمدت مطرح نمی‌شود.", size=10.5, color=MUTED, italic=True, before=8, after=0)

# 3 — Research fit
page_break(doc)
label(doc, "هم‌راستایی راهبردی")
h1(doc, "چرا این همکاری برای چارگون معنا دارد؟")
cards = [
    ("آسا و هوش مصنوعی محلی", "آسا بر تحلیل داده سازمانی، معماری ماژولار و حفظ داده داخل سازمان تأکید دارد. مارجین‌لیفت می‌تواند به‌صورت یک سرویس تخصصی کنار این معماری قرار گیرد.", True),
    ("هوش تجاری و تصمیم‌سازی", "هوش تجاری چارگون فاصله داده تا بینش را کم می‌کند. مارجین‌لیفت مرحله بعد را هدف می‌گیرد: تصمیم اقدام، هزینه مجاز و اندازه‌گیری نتیجه افزایشی.", False),
    ("تعامل‌پذیری", "رابط‌های مستند چارگون، اتصال امن سامانه‌های تراکنش، وفاداری یا ارتباط با مشتری را ممکن می‌کنند؛ بدون ورود دستی و بدون ساخت جزیره جدید.", False),
    ("انسان‌به‌انسان و بازارگاه", "تجربه انسانی‌تر یعنی تماس و مشوق کمتر اما مرتبط‌تر. چشم‌انداز همکاری با محصولات مکمل نیز مسیر عرضه مشترک را باز می‌گذارد.", False),
]
card_grid(doc, cards)
h2(doc, "نتیجه تحقیق")
para(doc, "چارگون همین امروز هوش مصنوعی، هوش تجاری و تعامل‌پذیری دارد؛ بنابراین جایگاه صحیح مارجین‌لیفت «رقیب تحلیلی» نیست. جایگاه درست، یک موتور تصمیم تخصصی و قابل‌سنجش برای مسائل درآمدی و نگهداشت مشتری است که می‌تواند در آینده به‌عنوان قابلیت مکمل در اکوسیستم چارگون عرضه شود.", size=11, bold=True, line=1.35)
callout(doc, "حساسیت مهم", "سایت رسمی چارگون، سامانه ارتباط با مشتری را به‌عنوان محصول مستقل معرفی نمی‌کند. از این‌رو پایلوت باید بر داده‌های تراکنش و تعامل مشتری نهایی متکی باشد و از مسیر تعامل‌پذیری به سامانه‌های موجود متصل شود.", fill="FFF5E7", accent=AMBER)

# 4 — Problem / distinction
page_break(doc)
label(doc, "فرصت محصولی")
h1(doc, "ریسک بالا، مجوز تخفیف نیست")
para(doc, "بسیاری از سامانه‌ها می‌توانند مشتری در معرض ریزش را پیش‌بینی کنند. اما یک مدیر درآمد باید به سؤال سخت‌تری پاسخ دهد: «کدام اقدام، برای کدام مشتری، واقعاً نتیجه را تغییر می‌دهد و آیا هزینه آن توجیه دارد؟»", size=11.2, after=10, line=1.35)
decision_table = doc.add_table(rows=5, cols=3)
decision_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(decision_table)
set_table_borders(decision_table, color=LINE, size=6, inside=True)
set_fixed_table_widths(decision_table, [Inches(2.15), Inches(2.05), Inches(2.45)])
headers = ["تصمیم پیشنهادی", "رفتار محتمل", "گروه مشتری"]
for c, t in enumerate(headers):
    cell = decision_table.cell(0, c)
    set_cell_shading(cell, INK)
    set_cell_margins(cell, top=130, start=140, bottom=130, end=140)
    p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run(t), size=9.2, color=WHITE, bold=True)
rows = [
    ("عدم اقدام؛ حفظ تجربه", "بدون مشوق می‌ماند", "خودبازگشت"),
    ("اقدام هدفمند", "فقط با اقدام تغییر می‌کند", "حساس به اقدام"),
    ("توقف هزینه و بررسی علت", "حتی با مشوق بازنمی‌گردد", "مقاوم"),
    ("محدودسازی تماس", "اقدام ممکن است بدتر کند", "واکنش منفی"),
]
for r, row in enumerate(rows, start=1):
    for c, t in enumerate(row):
        cell = decision_table.cell(r, c)
        set_cell_shading(cell, SOFT if r % 2 else WHITE)
        set_cell_margins(cell, top=125, start=135, bottom=125, end=135)
        p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(t), size=9.1, color=INK, bold=(c == 2))
para(doc, "", size=2, after=3)
h2(doc, "مارجین‌لیفت چه چیزی اضافه می‌کند؟")
add_bullet(doc, "ممیزی آمادگی داده: آیا داده برای تحلیل علّی آماده است یا فقط تشخیص تاریخی ممکن است؟")
add_bullet(doc, "ریسک و زمان‌بندی: چه کسی در معرض ریزش است و چه زمانی باید اقدام کرد؟")
add_bullet(doc, "اثر افزایشی و سیاست اقدام: اقدام برای چه کسی ارزش دارد و سقف هزینه چقدر است؟")
add_bullet(doc, "حلقه نتیجه: مقایسه پیش‌بینی با واقعیت و تصمیم توسعه، بازبینی یا توقف.")
callout(doc, "تمایز روشن", "مارجین‌لیفت جایگزین ارتباط با مشتری، مدیریت کمپین، آسا یا هوش تجاری نیست؛ لایه تصمیم اقتصادی و حلقه اثبات میان آن‌هاست.", fill=INK, accent=LIME)

# 5 — Solution + visual
page_break(doc)
label(doc, "راهکار پیشنهادی")
h1(doc, "فضای تصمیم؛ از فایل خام تا گزارش نهایی مدیر")
flow = doc.add_table(rows=1, cols=4)
flow.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(flow)
set_table_borders(flow, color=LINE, size=6, inside=True)
set_fixed_table_widths(flow, [Inches(1.66)] * 4)
steps = [
    ("۰۱", "ورودی امن", "شناسه ناشناس و رویداد"),
    ("۰۲", "تحلیل و سیاست", "ریسک، زمان و اقدام"),
    ("۰۳", "اجرای کنترل", "گروه کنترل و اقدام"),
    ("۰۴", "گزارش نتیجه", "اثر افزایشی و بازده"),
]
for c, (num, title, body) in enumerate(steps):
    cell = flow.cell(0, c)
    set_cell_shading(cell, INK if c == 2 else SOFT)
    set_cell_margins(cell, top=160, start=120, bottom=160, end=120)
    p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    set_run(p.add_run(num), size=8.5, color=LIME if c == 2 else GREEN, bold=True)
    para(cell, title, size=10.5, color=WHITE if c == 2 else INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(cell, body, size=8.4, color="DCE5E0" if c == 2 else MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.15)
para(doc, "", size=2, after=5)
if SCREEN.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SCREEN), width=Inches(6.25))
    p.paragraph_format.space_after = Pt(5)
para(doc, "نمونه واقعی از فضای نگهداشت مارجین‌لیفت؛ داده نمایشی است و هیچ نتیجه مشتری واقعی ادعا نمی‌شود.", size=8.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, italic=True)
callout(doc, "خروجی مدیر", "یک تصمیم روشن، عدد مالی با برچسب سطح شواهد، ریسک‌ها و اقدام بعدی؛ نه نمودارهای علمی بدون ترجمه اجرایی.", fill=SOFT, accent=GREEN)

# 6 — Pilot
page_break(doc)
label(doc, "طراحی پایلوت")
h1(doc, "مسیر ۶ تا ۱۰ هفته‌ای با ایستگاه‌های تصمیم روشن")
pilot = doc.add_table(rows=5, cols=4)
pilot.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(pilot)
set_table_borders(pilot, color=LINE, size=6, inside=True)
set_fixed_table_widths(pilot, [Inches(1.20), Inches(1.55), Inches(2.25), Inches(1.64)])
heads = ["تصمیم", "زمان", "خروجی", "مرحله"]
for c, t in enumerate(heads):
    cell = pilot.cell(0, c); set_cell_shading(cell, INK); set_cell_margins(cell)
    p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER); set_run(p.add_run(t), size=9, color=WHITE, bold=True)
pilot_rows = [
    ("ادامه / اصلاح داده", "۱ تا ۲ هفته", "آمادگی داده، خط پایه و قرارداد نتیجه", "آمادگی"),
    ("تأیید سیاست", "۱ تا ۲ هفته", "گروه هدف، اقدام، هزینه مجاز و طراحی کنترل", "طراحی"),
    ("ادامه / توقف موقت", "۲ تا ۶ هفته", "اجرای محدود، ثبت مواجهه و خط قرمزها", "اجرا"),
    ("توسعه / بازبینی / توقف", "۳ روز", "اثر افزایشی، بازده، ریسک و توصیه مدیریتی", "گزارش"),
]
for r, row in enumerate(pilot_rows, start=1):
    for c, t in enumerate(row):
        cell = pilot.cell(r, c); set_cell_shading(cell, SOFT if r % 2 else WHITE); set_cell_margins(cell)
        p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(t), size=8.9, color=GREEN if c == 0 else INK, bold=(c in [0, 3]))
para(doc, "", size=2, after=4)
h2(doc, "حداقل داده موردنیاز")
para(doc, "شناسه ناشناس مشتری، تاریخ رویداد یا تراکنش، محصول یا خدمت، درآمد یا حاشیه سود، هزینه اقدام یا مشوق، گروه کنترل و اقدام، ثبت مواجهه و نتیجه.", size=10.3, color=MUTED, line=1.3)
h2(doc, "قانون اعتماد")
add_bullet(doc, "بدون گروه کنترل سالم، خروجی فقط «برآورد تاریخی» است و ادعای علّی نمی‌شود.")
add_bullet(doc, "اگر نمونه، نسبت گروه‌ها یا پنجره نتیجه نامعتبر باشد، تصمیم به حالت نیازمند بازبینی می‌رود.")
add_bullet(doc, "هر عدد مالی با یکی از برچسب‌های برآورد تاریخی، برآورد پایلوت یا اثر افزایشی تأییدشده ارائه می‌شود.")
callout(doc, "تعریف موفقیت", "پایلوت موفق یعنی اثر افزایشی مثبت با خط قرمزهای سالم و امکان دفاع مالی؛ نه صرفاً دقت بالای مدل.", fill="FFF5E7", accent=AMBER)

# 7 — Success metrics
page_break(doc)
label(doc, "سنجش نتیجه")
h1(doc, "گزارشی که مدیر مالی و مدیر بازاریابی هر دو می‌فهمند")
metrics = doc.add_table(rows=6, cols=3)
metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(metrics)
set_table_borders(metrics, color=LINE, size=6, inside=True)
set_fixed_table_widths(metrics, [Inches(2.18), Inches(2.32), Inches(2.14)])
for c, t in enumerate(["قاعده تصمیم", "معنای مدیریتی", "شاخص"]):
    cell = metrics.cell(0, c); set_cell_shading(cell, INK); set_cell_margins(cell)
    p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER); set_run(p.add_run(t), size=9.1, color=WHITE, bold=True)
metric_rows = [
    ("مثبت و فراتر از آستانه توافق‌شده", "سود واقعی ناشی از اقدام", "سود افزایشی تأییدشده"),
    ("کمتر از baseline یا سقف بودجه", "بخشی از مشوق که حذف یا بازتخصیص می‌شود", "هزینه مشوق قابل حذف"),
    ("مثبت با ریسک کنترل‌شده", "بازده پایلوت پس از همه هزینه‌ها", "بازده سرمایه‌گذاری پایلوت"),
    ("در محدوده خط قرمز", "عدم آسیب به درآمد، رضایت یا گروه حساس", "درآمد در معرض ریسک"),
    ("قابل توضیح و قابل بازبینی", "فاصله مدل با واقعیت و نیاز به اصلاح", "شکاف پیش‌بینی/واقعیت"),
]
for r, row in enumerate(metric_rows, start=1):
    for c, t in enumerate(row):
        cell = metrics.cell(r, c); set_cell_shading(cell, SOFT if r % 2 else WHITE); set_cell_margins(cell)
        p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER); set_run(p.add_run(t), size=8.8, color=INK, bold=(c == 2))
para(doc, "", size=2, after=5)
callout(doc, "فرمول مالی", "سود افزایشی تأییدشده = درآمد افزایشی × حاشیه سود - هزینه مشوق - هزینه کانال - هزینه اجرای پایلوت", fill=INK, accent=LIME)
h2(doc, "خروجی نهایی")
para(doc, "یک گزارش مدیریتی فارسی شامل تصمیم اجرایی، عددهای مهم برای مدیر مالی، پیام برای مدیر بازاریابی و ارتباط با مشتری، محدودیت‌های شواهد، ریسک‌ها و سه اقدام بعدی. مدل زیر جدول و اصطلاحات علمی پنهان نمی‌شود؛ اما اتاق شواهد برای تیم داده در دسترس می‌ماند.", size=10.8, line=1.35)

# 8 — Architecture/security/roles
page_break(doc)
label(doc, "امنیت و اجرا")
h1(doc, "هم‌راستا با معماری اعتماد چارگون")
cards = [
    ("داده داخل سازمان", "استقرار در محل سازمان یا محیط اختصاصی؛ انتقال حداقلی و کنترل‌شده داده.", True),
    ("اطلاعات هویتی حداقلی", "شناسه ناشناس‌شده؛ بدون نیاز پیش‌فرض به نام، موبایل یا اطلاعات هویتی مستقیم.", False),
    ("کنترل دسترسی", "نقش‌ها، ثبت رخداد، نسخه سیاست و تاریخچه تصمیم برای ممیزی.", False),
    ("تصمیم محافظه‌کار", "عدم اقدام پیش‌فرض در نبود شواهد کافی یا سلامت پایین داده.", False),
]
card_grid(doc, cards)
h2(doc, "تقسیم مسئولیت در پایلوت")
roles = doc.add_table(rows=4, cols=3)
roles.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(roles)
set_table_borders(roles, color=LINE, size=6, inside=True)
set_fixed_table_widths(roles, [Inches(2.30), Inches(2.16), Inches(2.18)])
for c, t in enumerate(["مشتری پایلوت", "چارگون", "مارجین‌لیفت"]):
    cell = roles.cell(0, c); set_cell_shading(cell, INK); set_cell_margins(cell)
    p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER); set_run(p.add_run(t), size=9.2, color=WHITE, bold=True)
role_rows = [
    ("مالک داده و شاخص؛ تأیید قرارداد داده", "حامی محصول و هماهنگی دسترسی", "ممیزی داده و طراحی تحلیل"),
    ("اجرای کمپین و ثبت مواجهه", "هماهنگی تعامل‌پذیری، امنیت و استقرار", "طراحی سیاست و گروه کنترل"),
    ("تأیید هزینه، حاشیه سود و نتیجه", "بازبینی قابلیت عرضه مشترک", "گزارش، شواهد و توصیه توسعه یا توقف"),
]
for r, row in enumerate(role_rows, start=1):
    for c, t in enumerate(row):
        cell = roles.cell(r, c); set_cell_shading(cell, SOFT if r % 2 else WHITE); set_cell_margins(cell)
        p = cell.paragraphs[0]; rtl(p, WD_ALIGN_PARAGRAPH.CENTER); set_run(p.add_run(t), size=8.7, color=INK)
para(doc, "", size=2, after=4)
callout(doc, "پیش‌شرط حقوقی", "قرارداد پردازش داده، مالکیت خروجی، سطح خدمت، محرمانگی و مسئولیت تصمیم تجاری پیش از اجرای زنده نهایی می‌شوند.", fill="FFF5E7", accent=AMBER)

# 9 — Partnership and commercial
page_break(doc)
label(doc, "مدل همکاری")
h1(doc, "اول اثبات؛ سپس قرارداد مقیاس")
partnerships = [
    ("A | پایلوت هم‌برند", "مارجین‌لیفت اجرای تخصصی را بر عهده می‌گیرد؛ چارگون حامی محصول، مسیر فنی و دسترسی به مشتری مناسب را فراهم می‌کند. پیشنهاد شروع.", True),
    ("B | عرضه مشترک یا بازارگاه", "پس از اثبات نتیجه، بسته مشترک با سهم درآمد، SLA، پشتیبانی و ownership روشن عرضه می‌شود.", False),
    ("C | ماژول اختصاصی", "در صورت کشش بازار و تکرارپذیری، تجربه کاربری و استقرار عمیق‌تر با اکوسیستم چارگون یکپارچه می‌شود.", False),
]
table = doc.add_table(rows=3, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_rtl(table)
set_table_borders(table, color=LINE, size=6, inside=True)
set_fixed_table_widths(table, [Inches(6.64)])
for i, (title, body, best) in enumerate(partnerships):
    cell = table.cell(i, 0); set_cell_shading(cell, INK if best else WHITE); set_cell_margins(cell, top=170, start=200, bottom=170, end=200)
    p = cell.paragraphs[0]; rtl(p); set_run(p.add_run(title), size=12.2, color=LIME if best else GREEN, bold=True)
    para(cell, body, size=9.6, color="DCE5E0" if best else MUTED, after=0, line=1.25)
para(doc, "", size=2, after=5)
h2(doc, "چارچوب تجاری پیشنهادی")
add_bullet(doc, "ارزیابی: مبلغ ثابت برای بررسی آمادگی داده و تحلیل تاریخی؛ بدون ادعای علّی.")
add_bullet(doc, "پایلوت زنده: مبلغ ثابت برای طراحی گروه کنترل، اجرا و گزارش نتیجه.")
add_bullet(doc, "مقیاس: اشتراک یا سهم درآمد فقط بر مبنای صرفه‌جویی یا سود افزایشی تأییدشده و قابل ممیزی.")
callout(doc, "موضع پیشنهادی در مذاکره", "در جلسه اول قیمت نهایی ارائه نشود. ابتدا اندازه داده، چرخه خرید، هزینه اقدام، مدل استقرار و نقش چارگون در فروش/پشتیبانی روشن شود.", fill=SOFT, accent=GREEN)
para(doc, "این سند پیشنهاد اولیه و غیرالزام‌آور است. دامنه، قیمت، مسئولیت‌ها، مالکیت فکری و مفاد حقوقی پس از Discovery مشترک نهایی می‌شوند.", size=9, color=MUTED, italic=True, before=8, after=0)

# 10 — Next step + sources
page_break(doc)
label(doc, "قدم بعدی")
h1(doc, "یک جلسه؛ چهار تصمیم")
next_cards = [
    ("۱ | حامی سازمانی", "چه کسی در چارگون مالک بررسی محصول و تصمیم ادامه مسیر است؟", True),
    ("۲ | مشتری پایلوت", "کدام مشتری چرخه خرید تکرارشونده، داده کافی و درد مالی روشن دارد؟", False),
    ("۳ | معماری", "آیا انتقال امن دوره‌ای برای شروع کافی است یا اتصال مستقیم و استقرار محلی از روز اول لازم است؟", False),
    ("۴ | موفقیت", "چه شاخص و چه آستانه‌ای تصمیم توسعه، بازبینی یا توقف را تعیین می‌کند؟", False),
]
card_grid(doc, next_cards)
callout(doc, "درخواست مشخص", "تعیین یک جلسه ۶۰ دقیقه‌ای برای مرور دمو، قرارداد داده و انتخاب اولین پایلوت؛ بدون تعهد خرید یا توسعه اختصاصی.", fill=INK, accent=LIME)
h2(doc, "منابع عمومی مورد استفاده")
sources = [
    "سایت رسمی چارگون؛ محصولات، خدمات، مشتریان و معرفی شرکت: https://chargoon.com/",
    "پلتفرم هوش مصنوعی سازمانی آسا: https://chargoon.com/enterprise-ai-platform/",
    "هوش تجاری چارگون: https://chargoon.com/business-intelligence-software/",
    "تعامل‌پذیری و رابط‌های برنامه‌نویسی: https://chargoon.com/api-service/",
    "مرکز امنیت چارگون: https://chargoon.com/security-center/",
    "راهنمای برند چارگون: https://chargoon.com/logo/",
    "گفت‌وگوی مدیران چارگون درباره انسان‌به‌انسان، هوش مصنوعی محلی و بازارگاه: https://digiato.com/iran-technology-news/chargoon-interview-rebranding",
    "دموی مارجین‌لیفت: https://marginlift.ir/",
]
for src in sources:
    add_bullet(doc, src)
para(doc, "یادداشت دقت: در این سند هیچ ادعایی درباره دقت مدل، نتیجه مالی مشتری، تناسب محصول و بازار یا آمادگی استقرار در محیط عملیاتی مطرح نشده است. همه نتایج به پایلوت با داده واقعی نیاز دارند.", size=9.2, color=RED, bold=True, before=6, after=0, line=1.25)

doc.core_properties.title = "پیشنهاد همکاری MarginLift و چارگون"
doc.core_properties.subject = "پایلوت مشترک تصمیم‌یار نگهداشت مشتری"
doc.core_properties.author = "MarginLift"
doc.core_properties.keywords = "MarginLift, Chargoon, Retention, Uplift, Causal Decisioning"
doc.save(OUT)
print(OUT)
